from docx import Document
import sys
from mtuoctranslatefiles.abstract_file import AbstractFile


class Docx(AbstractFile):
    supported_file_extensions = ['.docx']

    def translate_paragraphs(self,paragraphs, translation_request):
        for paragraph in paragraphs:

            if len(paragraph.runs) == 1:
                paragraph.runs[0].text = translation_request(paragraph.runs[0].text)
            elif len(paragraph.runs) == 0:
                continue
            else:
                first_run_style = paragraph.runs[0].style
                #get the dominant run in the paragraph
                dominant_run = max(paragraph.runs, key=lambda x: len(x.text))
                dom_font = dominant_run.font.name
                dom_bold = dominant_run.font.bold
                dom_italic = dominant_run.font.italic
                dom_underline = dominant_run.font.underline

                paragraph.text = translation_request(paragraph.text)
                paragraph.runs[0].font.name = dom_font
                paragraph.runs[0].font.bold = dom_bold
                paragraph.runs[0].font.italic = dom_italic
                paragraph.runs[0].font.underline = dom_underline



    def translate(self, translation_request, file_path: str):
        outzip_path = self.get_output_path(file_path)
        document = Document(file_path)

        self.translate_paragraphs(document.paragraphs, translation_request)
        
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.translate_paragraphs(cell.paragraphs, translation_request)

        document.save(outzip_path)

        return outzip_path
